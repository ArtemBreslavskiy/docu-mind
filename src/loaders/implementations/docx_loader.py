import logging
from pathlib import Path
from tqdm import tqdm
from docx import Document as DocxDocument
from loaders.base import BaseLoader
from core_schemas import Document
from src.logger.logger_setup import get_logger


class DocxLoader(BaseLoader):
    def __init__(self, logger: logging.Logger = None, **kwargs):
        super().__init__(**kwargs)
        self.logger = logger or get_logger("pipeline")
        if logger is None:
            self.logger.disabled = True

    def load(self, raw_dir: str | Path, show_progress_bar: bool = True) -> list[Document]:
        raw_dir = Path(raw_dir)
        documents = []
        docx_files = list(raw_dir.glob("*.docx"))
        if not docx_files:
            self.logger.warning("No DOCX files found in %s", raw_dir)
            return []

        self.logger.info("Starting to load %d DOCX files from %s", len(docx_files), raw_dir)
        progress_bar = tqdm(docx_files, desc="Loading DOCX", unit="file") if show_progress_bar else docx_files

        for docx_file in progress_bar:
            try:
                doc = DocxDocument(docx_file)
                paragraphs = []
                for para in doc.paragraphs:
                    if para.style.name.startswith("Heading"):
                        paragraphs.append(f"## {para.text}")
                    else:
                        paragraphs.append(para.text)
                text = "\n".join(paragraphs)
                if not text.strip():
                    self.logger.debug("No text extracted from %s", docx_file)
                    continue

                props = doc.core_properties
                meta = {
                    "loaded_from": "docx",
                    "source": str(docx_file),
                    "title": props.title or docx_file.stem,
                    "author": props.author or "",
                }
                description = f"{meta['title']}--{meta.get('author', 'unknown')}"

                documents.append(Document(
                    content=text,
                    metadata=meta,
                    description=description
                ))
            except Exception as e:
                self.logger.error("Failed to process %s: %s", docx_file, e)

        self.logger.info("Finished loading DOCX files. Total documents: %d", len(documents))
        return documents
